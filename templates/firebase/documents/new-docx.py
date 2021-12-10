import urllib.request
import json
from docx import Document
from docx.shared import Inches, Pt

word_levels = {'w1':'初級', 'w2':'中級', 'w3':'中高級'}
column_names = ['級別', '單字', '意義']
url = 'http://localhost:5000/gept/gept-words_json/{}'

for wl in word_levels.keys():
    req = urllib.request.urlopen(url.format(wl))
    response_data = req.read()
    json_data = json.loads(response_data)

    # 設定一般字型
    document = Document()
    font = document.styles['Normal'].font
    font.size = Pt(12)

    # 檔案標題
    document.add_heading('GEPT 全民英檢字彙表', 0)
    p = document.add_paragraph()
    p.add_run('General English Proficiency Test Word Lists').italic = True
    document.add_heading(word_levels[wl]+'字彙', level=1)

    # 表格標題
    table = document.add_table(rows=1, cols=3, style="Light Grid")
    hdr_cells = table.rows[0].cells
    for i in range(3):
        hdr_cells[i].text = column_names[i]

    # 表格內容
    total = 0
    counter = 0
    for letter in json_data['jsonData']:
        # 每個級別新增 100 字彙
        if total == 100:
            break
        # 每個字母新增 10 個字彙
        counter = 0
        for word in letter['words']:
            if counter != 0 and counter % 10 == 0:
                break
            # 新增單字的長度大於等於 5     
            if len(word[1]) >= 5:
                total += 1
                counter += 1
                word[0] = str(total)
                print(word)
                row_cells = table.add_row().cells
                for i in range(3):
                    row_cells[i].text = word[i]

    document.save('gept-words-{}.docx'.format(wl))